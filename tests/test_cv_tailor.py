"""Tests pour backend/agents/cv_tailor.py (pipeline DOCX-template,
filter section-aware SUMMARY + Relevant coursework)."""
import json
from pathlib import Path
from unittest.mock import patch

import pytest

from backend.agents import cv_tailor, form_filler


SAMPLE_PROFILE = {
    "first_name": "Alex",
    "last_name": "Nitescu",
    "email": "alex@example.com",
    "phone": "+33 6 12 34 56 78",
    "city": "Paris",
    "country": "France",
    "linkedin": "https://linkedin.com/in/x",
    "github": "https://github.com/x",
    "current_title": "Data Scientist",
    "years_experience": 3,
    "skills": ["Python", "FastAPI"],
    "cv_path": "",
    "base_cv_path": "",
    "cv_output_dir": "",
}

SAMPLE_OFFER = {
    "title": "Senior Data Engineer F/H",
    "company": "BNP Paribas",
    "location": "Paris, France",
    "contract_type": "CDI",
    "skills": ["Python", "Spark", "AWS"],
    "missions": ["Build pipelines", "Optimize cost"],
}


@pytest.fixture
def profile_in_tmp(tmp_path, monkeypatch):
    """Patche le PROFILE_PATH partagé avec form_filler pour pointer dans tmp_path."""
    profile_path = tmp_path / "user_profile.json"
    profile_path.write_text(json.dumps(SAMPLE_PROFILE), encoding="utf-8")
    monkeypatch.setattr(form_filler, "PROFILE_PATH", profile_path)
    return profile_path


def write_profile(profile_path: Path, **overrides):
    data = {**SAMPLE_PROFILE, **overrides}
    profile_path.write_text(json.dumps(data), encoding="utf-8")


# ──────────────────────────────────────────────────────────────────────────
# slug / canonical title / filename
# ──────────────────────────────────────────────────────────────────────────


def test_slug_strips_accents_and_lowercases():
    assert cv_tailor._slug("L'Oréal") == "loreal"
    assert cv_tailor._slug("Société Générale") == "societe_generale"
    assert cv_tailor._slug("Welcome to the Jungle") == "welcome_to_the_jungle"
    assert cv_tailor._slug("Data/Scientist") == "data_scientist"
    assert cv_tailor._slug("  Hello   World  ") == "hello_world"


def test_slug_handles_none_and_empty():
    assert cv_tailor._slug(None) == ""
    assert cv_tailor._slug("") == ""
    assert cv_tailor._slug("---") == ""


def test_slug_allow_caps_preserves_case():
    assert cv_tailor._slug("BNP Paribas", allow_caps=True) == "BNP_Paribas"
    assert cv_tailor._slug("BNP Paribas") == "bnp_paribas"


def test_slug_title_uppercases_first_letter():
    assert cv_tailor._slug_title("alexandru nitescu") == "Alexandru_Nitescu"
    assert cv_tailor._slug_title("deep learning algorithm") == "Deep_Learning_Algorithm"


def test_slug_title_preserves_short_all_caps_acronyms():
    """BS, MS, AI, ML, NLP doivent rester ALL-CAPS, pas devenir Bs/Ms/Ai/Ml/Nlp."""
    assert cv_tailor._slug_title("BS MS Engineer") == "BS_MS_Engineer"
    assert cv_tailor._slug_title("AI/ML Engineer") == "AI_ML_Engineer"
    assert cv_tailor._slug_title("NLP Researcher") == "NLP_Researcher"


def test_slug_title_handles_none_and_empty():
    assert cv_tailor._slug_title(None) == ""
    assert cv_tailor._slug_title("") == ""


def test_canonical_job_title_strips_parentheticals():
    assert (
        cv_tailor._canonical_job_title(
            "Deep Learning Algorithm Graduate (TikTok Search Ranking)"
        )
        == "Deep Learning Algorithm Graduate"
    )


def test_canonical_job_title_strips_dash_suffix():
    assert (
        cv_tailor._canonical_job_title(
            "Deep Learning Algorithm Graduate - 2026 Start (BS/MS)"
        )
        == "Deep Learning Algorithm Graduate"
    )


def test_canonical_job_title_strips_em_and_en_dashes():
    assert cv_tailor._canonical_job_title("ML Engineer – 2026 New Grad") == "ML Engineer"
    assert cv_tailor._canonical_job_title("ML Engineer — Intern Track") == "ML Engineer"


def test_canonical_job_title_strips_gender_marker():
    assert cv_tailor._canonical_job_title("Senior Data Scientist F/H") == "Senior Data Scientist"
    assert cv_tailor._canonical_job_title("Data Engineer H/F") == "Data Engineer"
    assert cv_tailor._canonical_job_title("Lead ML M/F") == "Lead ML"


def test_canonical_job_title_preserves_slash_in_dual_skills():
    """Le strip F/H ne doit pas amputer AI/ML ou full-stack JS/TS."""
    assert cv_tailor._canonical_job_title("AI/ML Engineer") == "AI/ML Engineer"


def test_canonical_job_title_handles_none_and_empty():
    assert cv_tailor._canonical_job_title(None) == ""
    assert cv_tailor._canonical_job_title("") == ""


def test_make_filename_follows_new_convention():
    """0_CV_Firstname_Lastname_JobTitle.pdf — pas de company, Title case."""
    profile = {**SAMPLE_PROFILE, "first_name": "Alexandru", "last_name": "Nitescu"}
    offer = {"title": "Deep Learning Algorithm Graduate (TikTok Search Ranking) - 2026 Start (BS/MS)"}
    fname = cv_tailor.make_filename(profile, offer)
    assert fname == "0_CV_Alexandru_Nitescu_Deep_Learning_Algorithm_Graduate.pdf"


def test_make_filename_strips_gender_marker():
    offer = {"title": "Senior Data Engineer F/H"}
    fname = cv_tailor.make_filename(SAMPLE_PROFILE, offer)
    assert fname == "0_CV_Alex_Nitescu_Senior_Data_Engineer.pdf"


def test_make_filename_falls_back_when_fields_missing():
    profile = {"first_name": "Alex"}
    offer = {"title": "Dev"}
    fname = cv_tailor.make_filename(profile, offer)
    assert fname == "0_CV_Alex_X_Dev.pdf"


def test_resolve_output_path_uses_company_folder(tmp_path, profile_in_tmp):
    write_profile(profile_in_tmp, cv_output_dir=str(tmp_path / "CVs"))
    profile = form_filler.load_profile()
    path = cv_tailor.resolve_output_path(profile, SAMPLE_OFFER)
    assert path.parent.name == "BNP_Paribas"
    assert path.name.endswith(".pdf")
    assert path.parent.parent == tmp_path / "CVs"


def test_resolve_output_path_missing_dir_raises(profile_in_tmp):
    profile = form_filler.load_profile()
    with pytest.raises(ValueError, match="cv_output_dir"):
        cv_tailor.resolve_output_path(profile, SAMPLE_OFFER)


# ──────────────────────────────────────────────────────────────────────────
# DOCX reading helper (kept for diagnostics)
# ──────────────────────────────────────────────────────────────────────────


def test_read_base_cv_extracts_paragraphs(tmp_path):
    from docx import Document

    docx_path = tmp_path / "base.docx"
    doc = Document()
    doc.add_paragraph("Alex Nitescu")
    doc.add_paragraph("Data Scientist")
    doc.add_paragraph("")
    doc.add_paragraph("Skills: Python, FastAPI")
    doc.save(str(docx_path))

    text = cv_tailor.read_base_cv(docx_path)
    assert "Alex Nitescu" in text
    assert "Skills: Python, FastAPI" in text


def test_read_base_cv_missing_path_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        cv_tailor.read_base_cv(tmp_path / "nope.docx")


def test_read_base_cv_wrong_extension_raises(tmp_path):
    bad = tmp_path / "cv.pdf"
    bad.write_bytes(b"%PDF-fake")
    with pytest.raises(ValueError, match="docx"):
        cv_tailor.read_base_cv(bad)


# ──────────────────────────────────────────────────────────────────────────
# Substantive heuristic + section detection
# ──────────────────────────────────────────────────────────────────────────


def test_is_substantive_accepts_long_content():
    bullet = "Built and automated end-to-end data pipelines processing 100M+ records daily"
    assert cv_tailor._is_substantive(bullet) is True


def test_is_substantive_rejects_short_or_layout_text():
    assert cv_tailor._is_substantive("Alex Nitescu") is False
    assert cv_tailor._is_substantive("EXPERIENCE") is False
    assert cv_tailor._is_substantive("Exponens\tParis, France") is False
    assert cv_tailor._is_substantive("alex@example.com") is False


def test_is_section_header_accepts_all_caps_short_lines():
    assert cv_tailor._is_section_header("SUMMARY") is True
    assert cv_tailor._is_section_header("EXPERIENCE") is True
    assert cv_tailor._is_section_header("ADDITIONAL INFORMATION") is True
    # Avec tabulation de padding Word
    assert cv_tailor._is_section_header("SUMMARY\t") is True


def test_is_section_header_rejects_mixed_case_or_long_lines():
    assert cv_tailor._is_section_header("Summary") is False
    assert cv_tailor._is_section_header("Data Scientist") is False
    assert cv_tailor._is_section_header("THIS IS A VERY LONG ALL CAPS LINE THAT IS NOT A HEADER") is False


def test_normalize_section_maps_keywords():
    assert cv_tailor._normalize_section("SUMMARY") == "SUMMARY"
    assert cv_tailor._normalize_section("PROFILE") == "SUMMARY"
    assert cv_tailor._normalize_section("EXPERIENCE") == "EXPERIENCE"
    assert cv_tailor._normalize_section("EDUCATION") == "EDUCATION"
    assert cv_tailor._normalize_section("PROJECTS") == "PROJECTS"


def test_normalize_section_returns_none_for_non_header():
    assert cv_tailor._normalize_section("Built data pipelines processing 100M records") is None
    assert cv_tailor._normalize_section("Alex Nitescu") is None


# ──────────────────────────────────────────────────────────────────────────
# Section-aware editable filter
# ──────────────────────────────────────────────────────────────────────────


def _make_user_like_docx(path: Path) -> None:
    """DOCX mimant la structure d'un CV utilisateur réel : header, SUMMARY,
    EXPERIENCE (bullets), EDUCATION (avec Relevant coursework), PROJECTS."""
    from docx import Document

    doc = Document()
    # Header
    doc.add_paragraph("Alex Nitescu")
    doc.add_paragraph("2026 Data Scientist Graduate")
    doc.add_paragraph("Paris, France - alex@example.com")
    # SUMMARY
    doc.add_paragraph("SUMMARY")
    doc.add_paragraph(
        "Applied Mathematics graduate building production ML systems "
        "and full-stack AI products."
    )
    # EXPERIENCE
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Exponens\tParis, France")
    doc.add_paragraph(
        "Built and automated end-to-end data pipelines processing 100M+ records "
        "daily using Python, Pandas, PostgreSQL."
    )
    doc.add_paragraph(
        "Designed dynamic dashboards consumed by 5+ business teams to track "
        "revenue and operational KPIs."
    )
    # EDUCATION
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("CY Tech\tCergy, France")
    doc.add_paragraph(
        "Relevant coursework: Advanced Machine Learning, Deep Learning, NLP, "
        "LLMs, RAG, Time Series Analysis."
    )
    doc.add_paragraph(
        "LSTM-based Parkinson's disease classifier (TensorFlow/Keras) on voice "
        "features with 94% test accuracy."
    )
    doc.save(str(path))


def test_collect_editable_in_sections_picks_summary_content(tmp_path):
    """Le paragraphe descriptif sous SUMMARY est éditable."""
    docx_path = tmp_path / "cv.docx"
    _make_user_like_docx(docx_path)
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    texts = [t for _, t in editables]
    assert any("Applied Mathematics graduate" in t for t in texts)


def test_collect_editable_in_sections_picks_relevant_coursework(tmp_path):
    """La ligne 'Relevant coursework:' sous EDUCATION est éditable."""
    docx_path = tmp_path / "cv.docx"
    _make_user_like_docx(docx_path)
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    texts = [t for _, t in editables]
    assert any(t.lower().startswith("relevant coursework") for t in texts)


def test_collect_editable_ignores_experience_bullets(tmp_path):
    """Aucun bullet EXPERIENCE ne doit être proposé à Gemini, même substantiel."""
    docx_path = tmp_path / "cv.docx"
    _make_user_like_docx(docx_path)
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    texts = [t for _, t in editables]
    assert not any("Built and automated" in t for t in texts)
    assert not any("dashboards consumed" in t for t in texts)


def test_collect_editable_ignores_education_projects(tmp_path):
    """Les autres lignes sous EDUCATION (LSTM Parkinson, ...) restent gelées."""
    docx_path = tmp_path / "cv.docx"
    _make_user_like_docx(docx_path)
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    texts = [t for _, t in editables]
    assert not any("Parkinson" in t for t in texts)


def test_collect_editable_returns_empty_when_no_summary_or_coursework(tmp_path):
    from docx import Document

    docx_path = tmp_path / "minimal.docx"
    doc = Document()
    doc.add_paragraph("Alex Nitescu")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Some bullet long enough to be substantive content here.")
    doc.save(str(docx_path))

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    assert editables == []


# ──────────────────────────────────────────────────────────────────────────
# Paragraph walking + run-preserving replacement
# ──────────────────────────────────────────────────────────────────────────


def test_collect_paragraphs_walks_body_and_tables(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Body para A")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell L"
    table.cell(0, 1).text = "Cell R"

    paragraphs = cv_tailor._collect_paragraphs(doc)
    texts = [p.text for p in paragraphs]
    assert "Body para A" in texts
    assert "Cell L" in texts
    assert "Cell R" in texts


def test_set_paragraph_text_keeps_first_run_formatting(tmp_path):
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Original bold text")
    r.bold = True
    p.add_run(" plain trailing")

    cv_tailor._set_paragraph_text(p, "Tailored replacement content")

    assert p.runs[0].text == "Tailored replacement content"
    assert p.runs[0].bold is True
    assert all(r.text == "" for r in p.runs[1:])
    assert p.text == "Tailored replacement content"


# ──────────────────────────────────────────────────────────────────────────
# _parse_edits
# ──────────────────────────────────────────────────────────────────────────


def test_parse_edits_handles_json_object():
    edits = cv_tailor._parse_edits('{"0": "new 0", "3": "new 3"}')
    assert edits == {0: "new 0", 3: "new 3"}


def test_parse_edits_strips_code_fence():
    raw = '```json\n{"5": "tailored bullet"}\n```'
    assert cv_tailor._parse_edits(raw) == {5: "tailored bullet"}


def test_parse_edits_handles_empty_response():
    assert cv_tailor._parse_edits("") == {}
    assert cv_tailor._parse_edits("   \n  ") == {}


def test_parse_edits_raises_on_invalid_json():
    with pytest.raises(RuntimeError, match="JSON invalide"):
        cv_tailor._parse_edits("not a json")


def test_parse_edits_skips_non_integer_keys():
    edits = cv_tailor._parse_edits('{"2": "ok", "title": "should be ignored"}')
    assert edits == {2: "ok"}


# ──────────────────────────────────────────────────────────────────────────
# Availability
# ──────────────────────────────────────────────────────────────────────────


def test_is_available_requires_key_profile_and_paths(profile_in_tmp, monkeypatch):
    monkeypatch.setenv("GEMINI_API_KEY", "fake")
    assert cv_tailor.is_available() is False

    write_profile(profile_in_tmp, base_cv_path="/tmp/x.docx", cv_output_dir="/tmp/CVs")
    assert cv_tailor.is_available() is True

    monkeypatch.delenv("GEMINI_API_KEY")
    assert cv_tailor.is_available() is False


# ──────────────────────────────────────────────────────────────────────────
# Orchestrator (mocked Gemini + mocked PDF conversion)
# ──────────────────────────────────────────────────────────────────────────


def test_tailor_cv_orchestrates_end_to_end(tmp_path, profile_in_tmp, monkeypatch):
    docx_path = tmp_path / "base.docx"
    _make_user_like_docx(docx_path)

    out_dir = tmp_path / "CVs"
    write_profile(profile_in_tmp, base_cv_path=str(docx_path), cv_output_dir=str(out_dir))
    monkeypatch.setenv("GEMINI_API_KEY", "fake")

    # Look up real editable indices to forge the Gemini JSON response
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    assert len(editables) == 2  # SUMMARY content + Relevant coursework

    fake_json = json.dumps({
        str(editables[0][0]): "Tailored SUMMARY for BNP Paribas role",
        str(editables[1][0]): "Tailored coursework mentioning Spark and AWS",
    })

    with patch.object(cv_tailor, "_call_gemini", return_value=fake_json) as llm_mock, patch.object(
        cv_tailor, "_convert_docx_to_pdf"
    ) as pdf_mock:
        pdf_mock.side_effect = lambda src, dst: dst.write_bytes(b"%PDF-fake")
        result = cv_tailor.tailor_cv(SAMPLE_OFFER)

    assert llm_mock.called
    sent_prompt = llm_mock.call_args[0][0]
    assert "BNP Paribas" in sent_prompt
    assert "Applied Mathematics graduate" in sent_prompt  # SUMMARY content sent
    assert "Built and automated" not in sent_prompt  # EXPERIENCE bullets gelés

    pdf_mock.assert_called_once()
    assert result["edited_count"] == 2
    assert result["editable_count"] == 2
    assert result["filename"].endswith(".pdf")
    assert "BNP_Paribas" in result["folder"]


def test_tailor_cv_requires_base_cv_path(tmp_path, profile_in_tmp, monkeypatch):
    monkeypatch.setenv("GEMINI_API_KEY", "fake")
    write_profile(profile_in_tmp, cv_output_dir=str(tmp_path / "CVs"))
    with pytest.raises(ValueError, match="base_cv_path"):
        cv_tailor.tailor_cv(SAMPLE_OFFER)


def test_tailor_cv_requires_api_key(profile_in_tmp, monkeypatch):
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    with pytest.raises(RuntimeError, match="GEMINI_API_KEY"):
        cv_tailor.tailor_cv(SAMPLE_OFFER)


def test_tailor_cv_falls_back_when_no_editable_paragraphs(
    tmp_path, profile_in_tmp, monkeypatch
):
    """DOCX sans SUMMARY ni Relevant coursework : on convertit tel quel,
    pas d'erreur, pas d'appel Gemini."""
    from docx import Document

    docx_path = tmp_path / "base.docx"
    doc = Document()
    doc.add_paragraph("Alex Nitescu")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph(
        "Built data pipelines processing 100M records but no SUMMARY section here."
    )
    doc.save(str(docx_path))

    write_profile(
        profile_in_tmp,
        base_cv_path=str(docx_path),
        cv_output_dir=str(tmp_path / "CVs"),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "fake")

    with patch.object(cv_tailor, "_call_gemini") as llm_mock, patch.object(
        cv_tailor, "_convert_docx_to_pdf"
    ) as pdf_mock:
        pdf_mock.side_effect = lambda src, dst: dst.write_bytes(b"%PDF-fake")
        result = cv_tailor.tailor_cv(SAMPLE_OFFER)

    # Gemini n'est PAS appelé dans ce cas — rien à tailorer
    llm_mock.assert_not_called()
    pdf_mock.assert_called_once()
    assert result["editable_count"] == 0
    assert result["edited_count"] == 0


def test_tailor_cv_tolerates_empty_edits_dict(tmp_path, profile_in_tmp, monkeypatch):
    """Si Gemini renvoie {} (rien à changer), le pipeline ne crashe pas — il
    sauvegarde une copie identique du DOCX et la convertit."""
    docx_path = tmp_path / "base.docx"
    _make_user_like_docx(docx_path)

    write_profile(
        profile_in_tmp,
        base_cv_path=str(docx_path),
        cv_output_dir=str(tmp_path / "CVs"),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "fake")

    with patch.object(cv_tailor, "_call_gemini", return_value="{}"), patch.object(
        cv_tailor, "_convert_docx_to_pdf"
    ) as pdf_mock:
        pdf_mock.side_effect = lambda src, dst: dst.write_bytes(b"%PDF-fake")
        result = cv_tailor.tailor_cv(SAMPLE_OFFER)

    assert result["edited_count"] == 0
    assert result["editable_count"] == 2  # SUMMARY + Relevant coursework détectés


def test_tailor_cv_ignores_indices_outside_editable_set(
    tmp_path, profile_in_tmp, monkeypatch
):
    """Si Gemini hallucine un index hors de la liste editable, on l'ignore
    proprement sans toucher au paragraphe inattendu."""
    docx_path = tmp_path / "base.docx"
    _make_user_like_docx(docx_path)

    write_profile(
        profile_in_tmp,
        base_cv_path=str(docx_path),
        cv_output_dir=str(tmp_path / "CVs"),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "fake")

    # index 0 (header) n'est PAS dans la liste editable
    from docx import Document

    editables = cv_tailor._collect_editable_in_sections(Document(str(docx_path)))
    valid_idx = editables[0][0]
    rogue = json.dumps(
        {"0": "Hacked Name", str(valid_idx): "Tailored SUMMARY"}
    )

    with patch.object(cv_tailor, "_call_gemini", return_value=rogue), patch.object(
        cv_tailor, "_convert_docx_to_pdf"
    ) as pdf_mock:
        pdf_mock.side_effect = lambda src, dst: dst.write_bytes(b"%PDF-fake")
        result = cv_tailor.tailor_cv(SAMPLE_OFFER)

    assert result["edited_count"] == 1  # seule l'édition valide a passé

    # Vérifie que le nom (paragraphe index 0) n'a pas été touché
    saved = Document(result["saved_docx_path"])
    paras = cv_tailor._collect_paragraphs(saved)
    assert "Alex Nitescu" == paras[0].text


def test_banned_cliches_constant_is_non_empty():
    assert len(cv_tailor.BANNED_CLICHES) >= 5
    assert all(isinstance(p, str) and p == p.lower() for p in cv_tailor.BANNED_CLICHES)
