"""cover_letter — génère une lettre de motivation long-form via Gemini.

Pipeline :
    1. Charge le profil utilisateur (form_filler.load_profile).
    2. Appelle Gemini avec un prompt structuré (offer + profile + match
       optionnel). Le LLM écrit le corps de la lettre dans la langue de
       l'offre (détection autonome — instruction explicite dans le prompt).
    3. python-docx assemble un DOCX simple (en-tête contact + date +
       destinataire + corps + signature).
    4. pdf_convert.convert_docx_to_pdf produit le PDF dans le même dossier
       que le CV adapté, avec le préfixe `1_Cover_Letter_*`.

Garanties :
    - Banned clichés (mêmes que cv_tailor) → instruction dans le prompt,
      audit warn (pas raise) sur la sortie pour signaler une violation.
    - Pas d'invention factuelle (consigne explicite dans le prompt).
    - Match optionnel → orientation des "matched" / "missing" skills sans
      les inventer.
"""
from __future__ import annotations

import json
import logging
import os
from datetime import date
from pathlib import Path
from typing import Any

from .cv_tailor import (
    BANNED_CLICHES,
    _canonical_job_title,
    _slug,
    _slug_title,
)
from .form_filler import load_profile
from .pdf_convert import convert_docx_to_pdf

logger = logging.getLogger(__name__)

DEFAULT_MODEL = "gemini-2.5-flash"
MAX_PAYLOAD_CHARS = 16000

_BANNED_LIST = ", ".join(BANNED_CLICHES)

# Prompt en anglais : Gemini suit mieux les contraintes structurelles en
# anglais. La consigne de LANGUE EXPLICITE en fin force la sortie dans la
# langue de l'offre (FR/EN/etc.).
_SYSTEM = f"""You write a long-form, ATS-friendly cover letter for a job applicant.

You receive:
  - The OFFER (title, company, location, missions, required skills, summary).
  - The candidate PROFILE (name, current_title, experience, education,
    skills, summary).
  - Optional MATCH (matched_skills, missing_skills). When given, emphasise
    matched skills truthfully and NEVER claim a missing skill is present.

Output ONLY the body of the letter as plain text. No salutation block, no
contact header, no signature line — those are appended by the layout code.
Paragraphs separated by a blank line.

Structure (3-4 paragraphs, ~300-450 words):
  1. Opening — state the role you are applying for + a one-line hook tying
     your background to the company's mission or product.
  2. Why this role — connect 2-3 concrete profile facts (experience,
     projects, technologies) to specific missions or skills of the offer.
     Mirror offer keywords ONLY when supported by the profile.
  3. Why this company — brief, factual reason (mission, product, stack,
     market). If no factual hook, keep it short and pivot to the role.
  4. Closing — availability + invitation to continue the conversation.

EXPERIENCE INTERPRETATION — critical rule:
- `years_experience` alone does not say whether the experience was
  acquired post-graduation or during studies. ALWAYS read
  `experience_context` when present.
- Experience acquired DURING studies (apprenticeship/alternance,
  internships, work-study contracts) does NOT make the candidate a
  senior. A profile with 3 years of apprenticeship who just graduated
  is entry-level / new-grad / junior — frame the letter accordingly.
- For graduate / new-grad / entry-level / junior roles, present this
  apprenticeship experience as a STRENGTH (practical experience while
  studying) rather than pretending to be senior.

LANGUAGE RULE: write the entire letter in the SAME language as the OFFER
(title + summary). If the offer is in French, write in French. If in
English, in English. NEVER mix languages and NEVER translate the offer.

NEVER use these clichés (even once): {_BANNED_LIST}.
NEVER invent employers, dates, metrics, certifications, or technologies
that are not present in the PROFILE.
NO bullet points, NO headings, NO markdown — flowing paragraphs only."""


# ──────────────────────────────────────────────────────────────────────────
# Availability + Gemini call
# ──────────────────────────────────────────────────────────────────────────


def is_available() -> bool:
    """Vrai si la clé Gemini + le profil + cv_output_dir sont configurés."""
    if not os.getenv("GEMINI_API_KEY"):
        return False
    try:
        profile = load_profile()
    except FileNotFoundError:
        return False
    return bool(profile.get("cv_output_dir"))


def _call_gemini(prompt: str) -> str:
    """Appelle Gemini en mode texte brut. Isolé pour facilité de mock."""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    model = os.getenv("GEMINI_MODEL") or DEFAULT_MODEL

    logger.info("cover_letter: appel Gemini (%s)", model)
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="text/plain",
            temperature=0.4,
        ),
    )
    return response.text or ""


def _profile_for_llm(profile: dict[str, Any]) -> dict[str, Any]:
    """Sous-ensemble du profil utile pour la lettre. Exclut le PII non
    nécessaire (cv_path, base_cv_path) et la lettre type (pour éviter que
    Gemini la recopie au lieu de générer).

    `experience_context` est passé si présent — qualifie `years_experience`
    pour éviter qu'un junior fraîchement diplômé avec X ans d'alternance
    soit présenté comme senior dans la lettre (cf. règle dans `_SYSTEM`).
    """
    return {
        "first_name": profile.get("first_name"),
        "last_name": profile.get("last_name"),
        "current_title": profile.get("current_title"),
        "years_experience": profile.get("years_experience"),
        "experience_context": profile.get("experience_context"),
        "education_level": profile.get("education_level"),
        "education_field": profile.get("education_field"),
        "school": profile.get("school"),
        "city": profile.get("city"),
        "country": profile.get("country"),
        "skills": profile.get("skills"),
        "languages": profile.get("languages"),
        "summary": profile.get("summary"),
    }


def _build_prompt(
    offer: dict[str, Any],
    profile: dict[str, Any],
    match: dict[str, Any] | None = None,
) -> str:
    """Compose le prompt complet (system + offer + profile + match optionnel)."""
    offer_compact = {
        k: v
        for k, v in offer.items()
        if v and k not in ("description", "url", "from_cache", "application_id",
                           "seen_before", "application_status", "llm_used",
                           "llm_error", "match")
    }
    parts: list[str] = [
        _SYSTEM,
        f"--- OFFER ---\n{json.dumps(offer_compact, ensure_ascii=False, indent=2)}",
        f"--- PROFILE ---\n{json.dumps(_profile_for_llm(profile), ensure_ascii=False, indent=2)}",
    ]
    if match:
        match_block = {
            "matched_skills_emphasize_truthfully": list(match.get("matched_skills") or []),
            "missing_skills_do_not_claim_present": list(match.get("missing_skills") or []),
        }
        parts.append(
            "--- MATCH ---\n"
            + json.dumps(match_block, ensure_ascii=False, indent=2)
        )
    return "\n\n".join(parts)


# ──────────────────────────────────────────────────────────────────────────
# Generation
# ──────────────────────────────────────────────────────────────────────────


def generate_cover_letter(
    offer: dict[str, Any],
    profile: dict[str, Any],
    match: dict[str, Any] | None = None,
) -> str:
    """Renvoie le corps brut de la lettre (sans en-tête ni signature).

    Raises:
        RuntimeError: clé Gemini absente ou Gemini renvoie une chaîne vide.
    """
    if not os.getenv("GEMINI_API_KEY"):
        raise RuntimeError("GEMINI_API_KEY non configurée")

    prompt = _build_prompt(offer, profile, match)
    # Cap dur — un prompt très long fait exploser le budget gratuit du tier
    # gemini-2.5-flash. Profile + offer sont déjà sub-MAX_PAYLOAD_CHARS en
    # pratique, mais on garde un filet.
    if len(prompt) > MAX_PAYLOAD_CHARS:
        prompt = prompt[:MAX_PAYLOAD_CHARS]

    text = _call_gemini(prompt).strip()
    if not text:
        raise RuntimeError("Gemini a renvoyé un texte vide pour la lettre")

    # Audit clichés bannis : on warn (pas raise) — la consigne est déjà
    # dans le prompt, et sur un texte long il peut y avoir des matches
    # tangentiels (e.g. "passionate" dans un nom de produit). Le warn
    # permet à l'utilisateur de juger.
    lower = text.lower()
    for cliche in BANNED_CLICHES:
        if cliche in lower:
            logger.warning("cover_letter: cliché présent malgré la règle: %r", cliche)
    return text


# ──────────────────────────────────────────────────────────────────────────
# Filename + path resolution
# ──────────────────────────────────────────────────────────────────────────


def make_filename(profile: dict[str, Any], offer: dict[str, Any]) -> str:
    """Convention : `1_Cover_Letter_Firstname_Lastname_JobTitle.pdf`.

    Préfixe `1_` pour trier après le CV (`0_CV_*`). Réutilise la même
    logique de slug Title-case + canonicalisation du titre que le CV pour
    cohérence visuelle dans le dossier entreprise.
    """
    short_title = _canonical_job_title(offer.get("title"))
    parts = [
        "1",
        "Cover_Letter",
        _slug_title(profile.get("first_name")) or "X",
        _slug_title(profile.get("last_name")) or "X",
        _slug_title(short_title) or "Role",
    ]
    base = "_".join(p for p in parts if p)
    return f"{base}.pdf"


def resolve_output_path(profile: dict[str, Any], offer: dict[str, Any]) -> Path:
    """`{cv_output_dir}/{Company_Sanitized}/{filename}.pdf` — même dossier
    que le CV adapté, pour que l'utilisateur ait CV + lettre côte à côte."""
    root = profile.get("cv_output_dir", "").strip()
    if not root:
        raise ValueError("profile.cv_output_dir n'est pas configuré")
    company_dir = _slug(offer.get("company"), allow_caps=True) or "Unknown_Company"
    return Path(root).expanduser() / company_dir / make_filename(profile, offer)


# ──────────────────────────────────────────────────────────────────────────
# DOCX layout
# ──────────────────────────────────────────────────────────────────────────


def _build_docx(
    text: str,
    profile: dict[str, Any],
    offer: dict[str, Any],
    docx_path: Path,
) -> None:
    """Construit un DOCX minimaliste : en-tête contact + date +
    destinataire + corps + signature.

    Le style typo est volontairement sobre (police par défaut, taille 10
    sur le contact). On vise un rendu PDF propre, pas une mise en page
    élaborée — la lettre est avant tout du texte.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # En-tête : nom + contact
    name = " ".join(
        filter(None, [profile.get("first_name"), profile.get("last_name")])
    ) or "Candidate"
    header = doc.add_paragraph()
    run = header.add_run(name)
    run.bold = True
    run.font.size = Pt(13)

    contact_bits = [profile.get("email"), profile.get("phone"), profile.get("city")]
    contact = " · ".join(b for b in contact_bits if b)
    if contact:
        cp = doc.add_paragraph(contact)
        if cp.runs:
            cp.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Date + destinataire
    doc.add_paragraph(date.today().strftime("%d/%m/%Y"))
    company = (offer.get("company") or "").strip()
    location = (offer.get("location") or "").strip()
    if company:
        doc.add_paragraph(company)
    if location:
        doc.add_paragraph(location)
    doc.add_paragraph()  # spacer

    # Corps — split sur double-newline pour préserver les paragraphes
    for para in text.strip().split("\n\n"):
        cleaned = para.strip()
        if cleaned:
            doc.add_paragraph(cleaned)

    # Signature
    doc.add_paragraph()
    doc.add_paragraph(name)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


# ──────────────────────────────────────────────────────────────────────────
# Orchestrator
# ──────────────────────────────────────────────────────────────────────────


def tailor_cover_letter(
    offer: dict[str, Any],
    match: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Génère lettre + DOCX + PDF + sauvegarde dans le dossier entreprise.

    Returns:
        {
            "text": str,                # corps brut tel que renvoyé par Gemini
            "saved_path": str,          # PDF final
            "saved_docx_path": str,     # DOCX intermédiaire (modifiable manuellement)
            "filename": str,
            "folder": str,
        }

    Raises:
        FileNotFoundError: profil absent
        ValueError: cv_output_dir manquant
        RuntimeError: clé Gemini absente, Gemini renvoie vide, ou
            conversion PDF indisponible.
    """
    profile = load_profile()
    output_pdf = resolve_output_path(profile, offer)
    output_docx = output_pdf.with_suffix(".docx")

    text = generate_cover_letter(offer, profile, match)
    _build_docx(text, profile, offer, output_docx)
    convert_docx_to_pdf(output_docx, output_pdf)

    logger.info(
        "cover_letter: PDF généré -> %s (%d chars de corps)", output_pdf, len(text)
    )
    return {
        "text": text,
        "saved_path": str(output_pdf),
        "saved_docx_path": str(output_docx),
        "filename": output_pdf.name,
        "folder": str(output_pdf.parent),
    }
