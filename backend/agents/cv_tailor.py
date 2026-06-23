"""CV tailor — adapte le CV DOCX de base à une offre EN PRÉSERVANT sa forme.

Différence majeure avec une approche Markdown -> PDF : on n'invente PAS de
mise en page. Le DOCX de l'utilisateur est ouvert, on identifie les
paragraphes "éditables" (descriptions longues, bullets de réalisations,
summary), on demande à Gemini une version tailorée de leur texte, et on
réécrit le contenu textuel des runs en gardant le formatting (font, gras,
italique, alignement, espacement) tel qu'il est dans le DOCX. Le résultat
est un nouveau DOCX visuellement identique à l'original, sauf que les
bullets sont reformulés pour mirrorer l'offre. On convertit ensuite ce
DOCX en PDF via docx2pdf (Microsoft Word COM) ou LibreOffice headless en
fallback.

Pipeline :
    1. Charge profile.base_cv_path avec python-docx.
    2. Walk paragraphes top-level + paragraphes dans les cellules de table.
    3. Heuristique _is_editable filtre : assez long, pas de tab,
       pas en MAJUSCULES seules, pas d'info contact (email, téléphone, URL).
    4. Gemini reçoit (indices, texte original, offre, profil) → renvoie
       un JSON {idx: nouveau_texte}. On respecte la taille (±25%) et la
       structure (pas d'ajout/suppression de paragraphes).
    5. Pour chaque édition, on réécrit le run[0] et on vide les autres
       runs du paragraphe → formatting préservé, texte changé.
    6. Sauvegarde DOCX + conversion PDF dans
       {cv_output_dir}/{Company}/0_cv_*.{docx,pdf}.
"""
from __future__ import annotations

import json
import logging
import os
import re
import unicodedata
from pathlib import Path
from typing import Any

from .form_filler import load_profile
# Re-export sous l'ancien nom interne pour que les tests existants qui
# patchent `cv_tailor._convert_docx_to_pdf` continuent de marcher. Le
# converter vit maintenant dans pdf_convert pour pouvoir être réutilisé
# par cover_letter et tout autre agent produisant un PDF.
from .pdf_convert import convert_docx_to_pdf as _convert_docx_to_pdf  # noqa: F401

logger = logging.getLogger(__name__)

DEFAULT_MODEL = "gemini-2.5-flash"

# Clichés bannis dans le prompt. Conservés comme constante pour que les
# tests puissent vérifier qu'ils sont absents d'une sortie type.
BANNED_CLICHES: tuple[str, ...] = (
    "passionate",
    "driven",
    "hard-working",
    "hard working",
    "team player",
    "fast learner",
    "hit the ground running",
    "always on the lookout",
    "results-oriented",
    "results oriented",
    "self-motivated",
    "self motivated",
)

_SYSTEM = """You tailor a candidate's CV to a specific job offer by REWRITING
SPECIFIC TEXT PARAGRAPHS, preserving the candidate's exact CV layout.

You receive:
  - The job OFFER (title, company, required skills, missions, ...).
  - The candidate's structured PROFILE (contact, languages, etc.).
  - A numbered list of EDITABLE_PARAGRAPHS from the candidate's DOCX. Each
    entry is "{idx}: {original text}". These are the substantive paragraphs
    (summary, bullets, project descriptions). Other paragraphs (name,
    section headers, dates, locations, contact info) have ALREADY been
    excluded — never reference them, never produce keys for them.

For each paragraph, decide whether to rewrite or keep it:
  - Rewrite ONLY when you can make the bullet more relevant to the offer
    while staying strictly factual (use only facts contained in the
    original paragraph, the PROFILE, or the OFFER as keywords).
  - Stay within ±25 % of the original character count.
  - Keep one paragraph per index — DO NOT merge, split, add or remove
    bullets.
  - Use English action verbs. Mirror offer keywords when truthful.
  - NEVER invent metrics, employers, titles, technologies, or dates that
    are not in the original paragraph.

NEVER use these clichés (even once): passionate, driven, hard-working,
team player, fast learner, hit the ground running, always on the
lookout, results-oriented, self-motivated.

Output a STRICT JSON object. Keys are paragraph indices (as strings).
Values are the rewritten text. Omit a key if you choose to keep that
paragraph verbatim. No fences, no preamble — JSON only."""


# ──────────────────────────────────────────────────────────────────────────
# Availability + paths
# ──────────────────────────────────────────────────────────────────────────


def is_available() -> bool:
    """Vrai si la clé Gemini et le profil (avec base_cv_path + cv_output_dir)
    sont configurés."""
    if not os.getenv("GEMINI_API_KEY"):
        return False
    try:
        profile = load_profile()
    except FileNotFoundError:
        return False
    return bool(profile.get("base_cv_path")) and bool(profile.get("cv_output_dir"))


def read_base_cv(path: str | Path) -> str:
    """Extrait le texte brut d'un DOCX (paragraphes + cellules de table).

    Conservé pour diagnostics / tests. Le pipeline de tailoring n'utilise plus
    cette fonction — il manipule directement l'objet Document pour pouvoir
    préserver le formatting des runs lors du remplacement.
    """
    p = Path(path).expanduser()
    if not p.is_file():
        raise FileNotFoundError(f"base_cv_path introuvable: {p}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"base_cv_path doit être un .docx (reçu: {p.suffix})")

    from docx import Document

    doc = Document(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


# ──────────────────────────────────────────────────────────────────────────
# Sanitization + filename convention
# ──────────────────────────────────────────────────────────────────────────


def _slug(s: str | None, *, allow_caps: bool = False) -> str:
    """Normalise une chaîne pour un nom de fichier/dossier ASCII."""
    if not s:
        return ""
    text = unicodedata.normalize("NFKD", s)
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[\s/\\\-]+", "_", text)
    text = re.sub(r"[^A-Za-z0-9_]", "", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text if allow_caps else text.lower()


def _slug_title(s: str | None) -> str:
    """Slug + Title-case par token. Préserve les acronymes ALL-CAPS courts
    (BS, MS, AI, ML, NLP, ...) intacts au lieu de les rabattre en Title.
    Utilisé pour la convention de nommage CV_Firstname_Lastname_JobTitle."""
    if not s:
        return ""
    raw = _slug(s, allow_caps=True)
    tokens: list[str] = []
    for tok in raw.split("_"):
        if not tok:
            continue
        if tok.isupper() and len(tok) <= 4:
            tokens.append(tok)
        else:
            tokens.append(tok[:1].upper() + tok[1:].lower())
    return "_".join(tokens)


def _canonical_job_title(title: str | None) -> str:
    """Réduit un titre d'offre à sa forme canonique pour un nom de fichier.

    Drop : parenthétiques `(...)`, suffixes après tiret long/court/em,
    marqueurs de genre F/H, H/F, M/F.

        "Deep Learning Algorithm Graduate (TikTok Search Ranking) - 2026 Start (BS/MS)"
          -> "Deep Learning Algorithm Graduate"
        "Senior Data Scientist F/H"  -> "Senior Data Scientist"
        "AI/ML Engineer"             -> "AI/ML Engineer"  (slash préservé hors F/H)
    """
    if not title:
        return ""
    t = title
    t = re.sub(r"\s*\([^)]*\)", "", t)
    t = re.split(r"\s+[-–—]\s+", t, maxsplit=1)[0]
    t = re.sub(r"\s+[FfHhMm]/[FfHhMm]\b", "", t)
    return t.strip()


def make_filename(profile: dict[str, Any], offer: dict[str, Any]) -> str:
    """Convention : 0_CV_Firstname_Lastname_JobTitle.pdf

    - "CV" en majuscules, pas "cv".
    - Prénom / nom Title-case (Alexandru, Nitescu, pas alexandru_nitescu).
    - Titre d'offre canonicalisé (parenthétiques et suffixes drop, voir
      `_canonical_job_title`) et Title-case.
    - PAS de nom d'entreprise dans le filename — c'est déjà dans le dossier
      parent.
    """
    short_title = _canonical_job_title(offer.get("title"))
    parts = [
        "0",
        "CV",
        _slug_title(profile.get("first_name")) or "X",
        _slug_title(profile.get("last_name")) or "X",
        _slug_title(short_title) or "Role",
    ]
    base = "_".join(p for p in parts if p)
    return f"{base}.pdf"


def resolve_output_path(profile: dict[str, Any], offer: dict[str, Any]) -> Path:
    """{cv_output_dir}/{Company_Sanitized}/{filename}.pdf"""
    root = profile.get("cv_output_dir", "").strip()
    if not root:
        raise ValueError("profile.cv_output_dir n'est pas configuré")
    company_dir = _slug(offer.get("company"), allow_caps=True) or "Unknown_Company"
    return Path(root).expanduser() / company_dir / make_filename(profile, offer)


# ──────────────────────────────────────────────────────────────────────────
# DOCX walking + editable heuristic
# ──────────────────────────────────────────────────────────────────────────

_CONTACT_RE = re.compile(
    r"(\+\d{2,3}\s*\d|https?://|linkedin\.|github\.|\.com\b|\.fr\b|\bphone\b|\btel\b)",
    re.IGNORECASE,
)


def _is_substantive(text: str) -> bool:
    """Vrai si le texte est un paragraphe de contenu (vs mise en page).

    Règles cumulatives — toutes doivent être vraies :
      - longueur ≥ 25 caractères (filtre noms, dates courtes)
      - pas de tabulation (filtre les lignes alignées droite tab-séparées)
      - pas tout-en-MAJUSCULES (filtre les en-têtes de section)
      - pas d'@ (filtre les emails)
      - pas de pattern contact (téléphone, URL, github/linkedin)
    """
    s = text.strip()
    if len(s) < 25:
        return False
    if "\t" in s:
        return False
    letters = re.sub(r"[^A-Za-z]", "", s)
    if letters and letters == letters.upper() and len(s) < 80:
        return False
    if "@" in s:
        return False
    if _CONTACT_RE.search(s):
        return False
    return True


# Mapping mot-clé -> section canonique. On essaie d'absorber les variantes
# de naming usuelles (SUMMARY/PROFILE/OBJECTIVE, EDUCATION/ACADEMIC, etc.)
_SECTION_KEYWORDS: dict[str, tuple[str, ...]] = {
    "SUMMARY": ("SUMMARY", "PROFILE", "OBJECTIVE", "ABOUT"),
    "EXPERIENCE": ("EXPERIENCE", "WORK", "EMPLOYMENT"),
    "EDUCATION": ("EDUCATION", "ACADEMIC", "DIPLOMA"),
    "PROJECTS": ("PROJECT",),
    "SKILLS": ("SKILL", "TECH STACK"),
    "LANGUAGES": ("LANGUAGE",),
    "CERTIFICATIONS": ("CERTIFICATION", "AWARD", "CREDENTIAL", "ACCOMPLISHMENT"),
    "ADDITIONAL": ("ADDITIONAL", "OTHER", "MISC", "INTEREST"),
}


def _is_section_header(text: str) -> bool:
    """Vrai si le paragraphe ressemble à un en-tête de section (tout-en-CAPS,
    court). Strip whitespace + tabs avant l'analyse (Word ajoute parfois des
    tabulations de padding sur les headers)."""
    s = re.sub(r"\s+", " ", text).strip()
    if not s:
        return False
    letters = re.sub(r"[^A-Za-z]", "", s)
    if not letters:
        return False
    if letters != letters.upper():
        return False
    return len(s) <= 50


def _normalize_section(text: str) -> str | None:
    """Map un texte d'en-tête vers son tag de section canonique, ou None si
    ce n'est pas un en-tête (le caller ignore l'enregistrement)."""
    if not _is_section_header(text):
        return None
    t = text.strip().upper()
    for canonical, keywords in _SECTION_KEYWORDS.items():
        if any(kw in t for kw in keywords):
            return canonical
    # En-tête tout-caps non reconnu — on retourne quand même un tag (l'upper
    # du texte) pour qu'on n'édite plus rien sous cet en-tête inconnu.
    return t


def _collect_editable_in_sections(doc: Any) -> list[tuple[int, str]]:
    """Filtre section-aware : ne renvoie QUE les paragraphes à tailorer.

    Règle stricte (issue d'un retour utilisateur explicite) :
      - sous SUMMARY : tout paragraphe substantiel (1 ou 2 max en pratique)
      - sous EDUCATION : UNIQUEMENT la ligne "Relevant coursework..."
      - tout le reste (header, contact, EXPERIENCE bullets, PROJECTS desc,
        SKILLS, LANGUAGES, CERTIFICATIONS) reste gelé, formatting d'origine
        préservé.

    Returns:
        [(paragraph_index, original_text), ...] aligné avec
        `_collect_paragraphs(doc)`.
    """
    paragraphs = _collect_paragraphs(doc)
    current_section: str | None = None
    editables: list[tuple[int, str]] = []

    for i, p in enumerate(paragraphs):
        text = p.text
        if not text or not text.strip():
            continue

        new_section = _normalize_section(text)
        if new_section is not None:
            current_section = new_section
            continue

        if current_section == "SUMMARY" and _is_substantive(text):
            editables.append((i, text))
            continue

        if current_section == "EDUCATION" and text.lower().lstrip().startswith(
            "relevant coursework"
        ):
            editables.append((i, text))
            continue

    return editables


def _collect_paragraphs(doc: Any) -> list[Any]:
    """Walk top-level paragraphs + table cell paragraphs, dans l'ordre de
    lecture. Retourne la liste des objets Paragraph (python-docx).
    L'index dans cette liste sert de clé pour les éditions Gemini.
    """
    paras: list[Any] = []
    for p in doc.paragraphs:
        paras.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paras.append(p)
    return paras


def _set_paragraph_text(para: Any, new_text: str) -> None:
    """Remplace le texte d'un paragraphe en conservant le formatting de son
    premier run (font, gras, italique, couleur, taille). Les runs suivants
    sont vidés.

    Compromis assumé : un paragraphe avec runs de styles mixtes (e.g.
    moitié bold / moitié plain) perd cette mixité — l'ensemble prend le
    style du premier run. Sur les bullets d'un CV typique c'est invisible
    parce que le formatting est uniforme par paragraphe.
    """
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# ──────────────────────────────────────────────────────────────────────────
# Gemini call
# ──────────────────────────────────────────────────────────────────────────


def _call_gemini(prompt: str) -> str:
    """Appelle Gemini en mode JSON. Isolé pour faciliter le mock."""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    model = os.getenv("GEMINI_MODEL") or DEFAULT_MODEL

    logger.info("cv_tailor: appel Gemini (%s)", model)
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.4,
            response_mime_type="application/json",
        ),
    )
    return response.text or ""


def _build_prompt(
    editables: list[tuple[int, str]],
    offer: dict[str, Any],
    profile: dict[str, Any],
    match: dict[str, Any] | None = None,
) -> str:
    """Construit le prompt avec les paragraphes numérotés à tailorer.

    Si `match` est fourni (provenant de /match-score), on injecte un bloc
    MATCH qui liste les matched_skills (à mettre en valeur quand le profil
    les supporte déjà) et les missing_skills (à ne PAS inventer — on les
    indique pour que Gemini ne sur-promette pas dessus). Aucune injection
    si match=None : appel rétrocompatible.
    """
    offer_compact = {
        k: v
        for k, v in offer.items()
        if v and k not in ("description", "url")
    }
    paragraphs_block = "\n".join(f"{idx}: {text}" for idx, text in editables)
    parts = [
        _SYSTEM,
        f"--- OFFER ---\n{json.dumps(offer_compact, ensure_ascii=False, indent=2)}",
        f"--- PROFILE ---\n{json.dumps(profile, ensure_ascii=False, indent=2)}",
    ]
    if match:
        matched = list(match.get("matched_skills") or [])
        missing = list(match.get("missing_skills") or [])
        match_block = {
            "matched_skills_emphasize_truthfully": matched,
            "missing_skills_do_not_claim_present": missing,
        }
        parts.append(
            "--- MATCH ---\n"
            "Use this to orient phrasing. Mirror matched skills when the "
            "original paragraph already supports them. Never claim a missing "
            "skill is present.\n"
            f"{json.dumps(match_block, ensure_ascii=False, indent=2)}"
        )
    parts.append(f"--- EDITABLE_PARAGRAPHS ---\n{paragraphs_block}")
    return "\n\n".join(parts) + "\n"


def _parse_edits(raw: str) -> dict[int, str]:
    """Parse la réponse JSON de Gemini en {idx: nouveau_texte}.

    Tolère un wrapper code-fence si Gemini en remet malgré le mime-type.
    """
    cleaned = raw.strip()
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", cleaned, re.DOTALL)
    if fence:
        cleaned = fence.group(1).strip()
    if not cleaned:
        return {}
    try:
        payload = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Gemini JSON invalide: {exc}") from exc
    if not isinstance(payload, dict):
        raise RuntimeError("Gemini n'a pas renvoyé un objet JSON")
    edits: dict[int, str] = {}
    for k, v in payload.items():
        try:
            edits[int(k)] = str(v).strip()
        except (TypeError, ValueError):
            logger.warning("cv_tailor: clé non-entière ignorée — %r", k)
    return edits


# ──────────────────────────────────────────────────────────────────────────
# Orchestrator
# ──────────────────────────────────────────────────────────────────────────


def tailor_cv(
    offer: dict[str, Any],
    match: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Édite en place les paragraphes éditables du DOCX base et exporte en PDF.

    Args:
        offer: champs structurés de l'offre (title, company, skills, ...).
        match: optionnel — {matched_skills, missing_skills} de /match-score.
            Injecté dans le prompt pour orienter Gemini sans inventer.

    Returns:
        {
            "saved_path": str,           # chemin absolu du PDF final
            "saved_docx_path": str,      # chemin absolu du DOCX tailoré
            "filename": str,             # nom du PDF seul
            "folder": str,               # dossier entreprise
            "edited_count": int,         # nb de paragraphes effectivement modifiés
            "editable_count": int,       # nb de paragraphes envoyés au LLM
        }

    Raises:
        FileNotFoundError: profil ou base_cv_path absent
        ValueError: cv_output_dir manquant ou DOCX invalide
        RuntimeError: clé API absente, Gemini KO, ou conversion PDF impossible
    """
    if not os.getenv("GEMINI_API_KEY"):
        raise RuntimeError("GEMINI_API_KEY non configurée")

    profile = load_profile()
    base_path = profile.get("base_cv_path", "").strip()
    if not base_path:
        raise ValueError("profile.base_cv_path n'est pas configuré (DOCX source manquant)")

    base = Path(base_path).expanduser()
    if not base.is_file():
        raise FileNotFoundError(f"base_cv_path introuvable: {base}")
    if base.suffix.lower() != ".docx":
        raise ValueError(f"base_cv_path doit être un .docx (reçu: {base.suffix})")

    output_pdf = resolve_output_path(profile, offer)
    output_docx = output_pdf.with_suffix(".docx")

    # 1. Charger le DOCX
    from docx import Document

    doc = Document(str(base))
    paragraphs = _collect_paragraphs(doc)

    # 2. Identifier les paragraphes éditables (section-aware, stricte) :
    #    SUMMARY content + 'Relevant coursework...' dans EDUCATION. Le reste
    #    du CV (EXPERIENCE bullets, PROJECTS, SKILLS, LANGUAGES, headers,
    #    contact) reste gelé. C'est volontairement conservateur — l'objectif
    #    est de ne JAMAIS toucher au formatting du CV de base, et de cibler
    #    uniquement les deux endroits qui ont un sens à être tailorés à
    #    l'offre.
    editables = _collect_editable_in_sections(doc)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    if not editables:
        # Pas d'erreur — on convertit le DOCX tel quel en PDF. L'utilisateur
        # récupère son CV original dans le dossier de l'offre.
        logger.warning(
            "cv_tailor: aucun paragraphe SUMMARY/coursework détecté, "
            "conversion DOCX -> PDF sans tailoring"
        )
        doc.save(str(output_docx))
        _convert_docx_to_pdf(output_docx, output_pdf)
        return {
            "saved_path": str(output_pdf),
            "saved_docx_path": str(output_docx),
            "filename": output_pdf.name,
            "folder": str(output_pdf.parent),
            "edited_count": 0,
            "editable_count": 0,
        }
    logger.info("cv_tailor: %d paragraphes éditables détectés", len(editables))

    # 3. Demander à Gemini les versions tailorées
    prompt = _build_prompt(editables, offer, profile, match)
    raw = _call_gemini(prompt)
    edits = _parse_edits(raw)
    logger.info("cv_tailor: %d éditions reçues sur %d demandées", len(edits), len(editables))

    # 4. Appliquer les éditions, en gardant l'index original
    editable_idxs = {idx for idx, _ in editables}
    edited_count = 0
    for idx, new_text in edits.items():
        if idx not in editable_idxs:
            logger.warning("cv_tailor: idx %d hors liste editable, ignoré", idx)
            continue
        if not new_text:
            continue
        _set_paragraph_text(paragraphs[idx], new_text)
        edited_count += 1

    # 5. Sauvegarder DOCX puis convertir en PDF
    doc.save(str(output_docx))
    _convert_docx_to_pdf(output_docx, output_pdf)

    logger.info(
        "cv_tailor: PDF généré -> %s (%d/%d éditions appliquées)",
        output_pdf,
        edited_count,
        len(editables),
    )
    return {
        "saved_path": str(output_pdf),
        "saved_docx_path": str(output_docx),
        "filename": output_pdf.name,
        "folder": str(output_pdf.parent),
        "edited_count": edited_count,
        "editable_count": len(editables),
    }
